# Import the libraries
import os
from shutil import copytree
from shutil import rmtree
from docx import Document
from datetime import date
from docx.shared import Pt
import pandas as pd
import re
import bs4
import requests
import random

JOB_TITLE_CLASS = '.icl-u-xs-mb--xs.icl-u-xs-mt--none.jobsearch-JobInfoHeader-title'
COMPANY_NAME_CLASS = '.icl-u-lg-mr--sm.icl-u-xs-mr--xs'
LISTING_CLASS = '#viewJobSSRRoot'
r_document = None


# Method from https://stackoverflow.com/questions/24805671/how-to-use-python-docx-to-replace-text-in-a-word-document
# -and-save
def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def replace_items(items):
    for index, item in enumerate(items):
        replace = r"" + item
        regex = re.compile(r"__spot" + str(index+1) + "__")
        docx_replace_regex(r_document, regex, replace)


def application():
    global r_document

    # List of all possible technical expertise
    expertise_list = list(pd.read_csv('technical_expertise.csv'))

    # Prompt the user for the indeed job listing url
    url = input("Enter url of job listing: ")
    if url == 'cancel':
        quit()
    
    # Create the request and get the soup
    res = requests.get(url)
    soup = bs4.BeautifulSoup(res.text, 'lxml')

    # Get the and print the position and company for the listing
    position = soup.select_one(JOB_TITLE_CLASS).text
    company = soup.select_one(COMPANY_NAME_CLASS).text
    print(f'Position is: {position}')
    print(f'Company is: {company}')

    # Display message to user saying application is created
    print(f'Creating application for position: {position} at company: {company}')

    # Set up the paths
    PATHS = {'src': '.template',
             'c': company + '/Cover Letter - ' + company + '.docx',
             'l': company + '/Listing - ' + company + '.txt',
             'r': company + '/Resume - ' + company + '.docx'}

    try:
        # Copy the template directory to a new name and rename it to the company name
        copytree(PATHS['src'], company)
    except:
        if not input('An application already exists for this company. Do you want to replace it? (Y or N)').lower() == 'y':
            return
        else:
            rmtree(company)
            print('Removed existing application')
            copytree(PATHS['src'], company)


    # Rename listing, resume, and cover letter in the newly created directory to include company name
    os.rename(company + '/Listing -.txt', PATHS['l'])
    os.rename(company + '/Resume -.docx', PATHS['r'])
    os.rename(company + '/Cover Letter -.docx', PATHS['c'])

    # Get todays date in correct format for cover letter
    today = date.today().strftime("%B %d, %Y")

    # Create document object for cover letter
    cl_document = Document(PATHS['c'])

    # Create correct font for the cover letter document
    c_style = cl_document.styles['Normal']
    c_font = c_style.font
    c_font.name = 'Tahoma'
    c_font.size = Pt(10)

    # Loop through each paragraph in the cover letter
    for paragraph in cl_document.paragraphs:
        # Find and replace __company__ in document to company name
        if '__company__' in paragraph.text:
            paragraph.text = paragraph.text.replace('__company__', company)
            paragraph.style = cl_document.styles['Normal']
        # Find and replace __position in document to position name
        if '__position__' in paragraph.text:
            paragraph.text = paragraph.text.replace('__position__', position)
            paragraph.style = cl_document.styles['Normal']
        # Find and replace __date__ in the document with formatted date
        if '__date__' in paragraph.text:
            paragraph.text = paragraph.text.replace('__date__', today)
            paragraph.style = cl_document.styles['Normal']

    # Save the updated cover letter
    cl_document.save(PATHS['c'])

    # Create document object for resume
    r_document = Document(PATHS['r'])

    # List of the technical expertise matches from the listing
    matches = []

    # Match listing items with the expertise items and add them to matches list
    print(f'MATCHES: ', end='')
    for item in expertise_list:
        pattern = r'\W' + item + r'\W'
        text = soup.select_one(LISTING_CLASS).get_text('\n')
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            match_text = match.group(0)
            matches.append(item)
            print(item, end=' ')
    print()
    items = list(set(matches))

    print('FILLS: ', end='')
    index = 0

    # Fill remaining expertise slots in order until there is 16
    while len(items) < 16:
        item = expertise_list[index]
        if item not in items:
            items.append(item)
            print(item, end=' ')
        index += 1
    print()

    # Updates the resume with the relevant items
    replace_items(items)

    # Save the resume document
    r_document.save(PATHS['r'])

    # Update the listing document, from user Abhishek Divekar:
    # https://stackoverflow.com/questions/36039919/beautifulsoup-output-to-txt-file
    listing_file = open(PATHS['l'], 'w', encoding='utf-8')
    listing_file.write(soup.select_one(LISTING_CLASS).get_text('\n'))
    listing_file.close()


if __name__ == '__main__':
    while True:
        application()
        if not input('Finished creating application, would you like to create another one? (Y or N)').lower() == 'y':
            print('breaking')
            break

    print("Exiting")
